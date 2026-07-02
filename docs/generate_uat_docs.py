#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""توليد دليل اختبار الأدوار: Markdown + Word + PDF."""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path

DOCS = Path(__file__).resolve().parent
SOURCE_PATH = DOCS / "_uat_source.md"
MD_PATH = DOCS / "اختبار-الأدوار-والصلاحيات.md"
DOCX_PATH = DOCS / "اختبار-الأدوار-والصلاحيات.docx"
PDF_PATH = DOCS / "اختبار-الأدوار-والصلاحيات.pdf"

PRIMARY = "5B2EFF"
HEADER_BG = "EDE9FF"
ALT_ROW = "F8F7FC"
ACCENT = "7C5CFF"
CHECKBOX = "☐"
BODY_FONT_PT = 10
APPENDIX_FONT_PT = 15  # +50% عن النص العادي
COVER_TITLE_PT = 28
COVER_SUBTITLE_PT = 16
COVER_META_PT = 13
CHECKBOX_FONT_PT = 17
CHECKBOX_ENV_FONT_PT = 26  # +150% عن مربع الصفوف العادية
REPORT_DEFECT_COUNT = 10
TABLE_HDR_CHECK = ""  # لا مربع تصحيح قصاد عناوين الجدول


def issue_date_str() -> str:
    return date.today().strftime("%d-%m-%Y")


def clean_uat_text(text: str) -> str:
    """إزالة المسارات والصلاحيات التقنية — الإبقاء على اسم الميزة فقط."""
    text = re.sub(r"\*\*", "", text)
    text = re.sub(r"\s*\(يتطلب\s*`?perm_\w+`?\)", "", text, flags=re.I)
    text = re.sub(r"\s*\(إن مُنح\s*`?perm_\w+`?\)", "", text, flags=re.I)
    text = re.sub(r"\s*—?\s*إن مُنح(?:ت)?\s*`?perm_\w+`?", "", text, flags=re.I)
    text = re.sub(r"مع\s*`?perm_\w+`?", "عند منح الصلاحية", text, flags=re.I)
    text = re.sub(r"بدون\s*`?perm_\w+`?", "بدون الصلاحية", text, flags=re.I)
    text = re.sub(r"`?perm_\w+`?", "", text, flags=re.I)
    text = re.sub(r"`/[^`]*`", "", text)
    text = re.sub(r"\(\s*أو\s+", "(", text)
    text = re.sub(r"\(\s*\)", "", text)
    text = re.sub(r"\s+من\s+`/login`", "", text)
    text = re.sub(r"\s*→\s*[^→]*`/[^`]+`", " → مرفوض", text)
    text = re.sub(r"\s{2,}", " ", text)
    text = re.sub(r"\s+([،.؛:])", r"\1", text)
    return text.strip(" —-")


def report_template_lines() -> list[str]:
    lines = [
        "الدور: _______________",
        "المختبِر: _______________",
        "التاريخ: _______________",
        "البيئة: □ تطوير  □ إنتاج",
        "",
        "عدد النقاط المنفّذة: ___ / ___",
        "عدد النقاط الناجحة:  ___ / ___",
        "عدد العيوب المكتشفة: ___",
        "",
        f"أهم {REPORT_DEFECT_COUNT} عيوب:",
    ]
    lines.extend(f"{i}. " for i in range(1, REPORT_DEFECT_COUNT + 1))
    lines.extend(
        [
            "",
            "النتيجة النهائية: □ ناجح  □ ناجح بتحفظات  □ فاشل",
            "",
            "ملاحظات:",
            "",
        ]
    )
    return lines


def _parse_test_items(tests_block: str) -> list[dict]:
    sections: list[dict] = []
    current_cat = "عام"
    items: list[str] = []
    table_row_new = re.compile(r"^\|\s*☐\s*\|\s*(.+?)\s*\|\s*(\d+)\s*\|$")
    table_row_old = re.compile(r"^\|\s*(\d+)\s*\|\s*(.+?)\s*\|\s*☐\s*\|$")

    def flush():
        nonlocal items
        if items:
            sections.append({"category": current_cat, "items": items})
            items = []

    for raw in tests_block.splitlines():
        line = raw.strip()
        if line.startswith("### "):
            flush()
            current_cat = clean_uat_text(line[4:].strip())
        elif line.startswith("- [ ]"):
            item = clean_uat_text(line[5:].strip())
            items.append(item)
        elif line.startswith("  - [ ]"):
            item = clean_uat_text(line[7:].strip())
            items.append(f"↳ {item}")
        elif m := table_row_new.match(line):
            items.append(clean_uat_text(m.group(1).strip()))
        elif m := table_row_old.match(line):
            items.append(clean_uat_text(m.group(2).strip()))
        elif line.startswith("|---"):
            continue
        elif line.startswith("| # |"):
            continue
    flush()
    return sections


def parse_markdown_roles(text: str) -> list[dict]:
    roles = []
    blocks = re.split(r"\n---\n", text)
    role_pattern = re.compile(
        r"# (\d+)\. (.+?) \((.+?)\)(?:\s*—[^\n]*)?\s*\n\n"
        r"## \d+\.1 نبذة عن الدور\s*\n\n(.+?)\n\n"
        r"## \d+\.2 نقاط الاختبار — الصلاحيات\s*\n(.*?)\n\n"
        r"## \d+\.3 التقرير المتوقع بعد الاختبار\s*\n\n(.+?)(?=\n\n\*\*معايير النجاح|\Z)",
        re.S,
    )
    for block in blocks:
        m = role_pattern.search(block)
        if not m:
            continue
        num, title, key, intro, tests_block, report_block = m.groups()
        key = key.strip()
        sections = _parse_test_items(tests_block)

        report_lines = []
        for line in report_block.splitlines():
            line = line.strip()
            if line.startswith("- "):
                report_lines.append(re.sub(r"\*\*", "", line[2:].strip()))
        success = ""
        sm = re.search(r"\*\*معايير النجاح:\*\* (.+)", block)
        if sm:
            success = sm.group(1).strip()

        roles.append(
            {
                "num": int(num),
                "title": title.strip(),
                "key": key.strip(),
                "intro": intro.strip(),
                "sections": sections,
                "report": report_lines,
                "success": success,
            }
        )
    return roles


def table_md(rows: list[str], start: int = 1) -> tuple[str, int]:
    lines = [
        "",
        f"|  | الميزة / نقطة الاختبار | # |",
        "|:---:|:---|---:|",
    ]
    n = start
    for row in rows:
        safe = row.replace("|", "\\|")
        lines.append(f"| {CHECKBOX} | {safe} | {n} |")
        n += 1
    lines.append("")
    return "\n".join(lines), n


def build_markdown(roles: list[dict]) -> str:
    parts = [
        "# دليل اختبار الأدوار والصلاحيات — نظام مرحّاب",
        "",
        "> **الغرض:** ملف اختبار يدوي (UAT) — لكل دور: نبذة → جدول صلاحيات → تقرير متوقع.",
        "",
        "## كيفية الاستخدام",
        "",
        "1. سجّل دخول بحساب الدور.",
        "2. نفّذ كل صف في الجدول.",
        "3. ضع ✓ في عمود التصحيح عند النجاح.",
        "4. راجع قسم «التقرير المتوقع».",
        "",
        "| الدور | مسار الدخول |",
        "|-------|-------------|",
        "| مدير النظام | `/login` → `/dashboard` |",
        "| مدير المنصة | `/login` → `/platform/dashboard` |",
        "| مدير الفعالية | `/login` → `/event-manager/dashboard` |",
        "| منظم الفعالية | `/login` → `/event-organizer/dashboard` |",
        "| المنسّق | `/login` → `/coordinator/check-in` |",
        "| مدير الدخول | `/login` → `/entry-manager/check-in` |",
        "| الضيف | `/i/{token}` — بدون تسجيل دخول |",
        "",
        "---",
        "",
    ]
    for role in roles:
        parts.append(f"# {role['num']}. {role['title']} ({role['key']})")
        parts.append("")
        parts.append(f"## {role['num']}.1 نبذة عن الدور")
        parts.append("")
        parts.append(role["intro"])
        parts.append("")
        parts.append(f"## {role['num']}.2 نقاط الاختبار — الصلاحيات")
        parts.append("")
        counter = 1
        for sec in role["sections"]:
            parts.append(f"### {sec['category']}")
            tbl, counter = table_md(sec["items"], counter)
            parts.append(tbl)
        parts.append(f"## {role['num']}.3 التقرير المتوقع بعد الاختبار")
        parts.append("")
        parts.append("بعد اكتمال الاختبار، يجب أن يكون المستخدم قادراً على:")
        parts.append("")
        for r in role["report"]:
            parts.append(f"- {r}")
        parts.append("")
        if role["success"]:
            parts.append(f"**معايير النجاح:** {role['success']}")
        parts.append("")
        parts.append("---")
        parts.append("")

    e2e_steps = [
        ("إنشاء منصة", "مدير النظام", "منصة + حساب مدير منصة"),
        ("إنشاء فعالية + أقسام + جدول", "مدير المنصة / مدير الفعالية", "فعالية جاهزة"),
        ("إضافة ضيوف وإرسال دعوات", "مدير الفعالية", "روابط دعوة تصل للضيوف"),
        ("تأكيد حضور ضيف", "الضيف", "QR يظهر"),
        ("إضافة منسّق + مدير دخول", "مدير الفعالية", "حسابات جاهزة"),
        ("بدء الفعالية", "مدير الفعالية", "«تعمل الآن»"),
        ("مسح QR عند البوابة", "مدير الدخول", "«حضر»"),
        ("مسح QR عند الطاولة", "المنسّق", "«جلس»"),
        ("تشغيل بث وإرسال الرابط", "منظم الفعالية", "مشاهدة البث من الدعوة"),
        ("إنهاء الفعالية", "مدير الفعالية", "«منتهية»"),
    ]
    parts.extend(["# ملحق: سيناريو End-to-End", ""])
    parts.append("|  | الخطوة | الدور | النتيجة المتوقعة | # |")
    parts.append("|:---:|--------|-------|------------------|---:|")
    for i, (step, role_name, result) in enumerate(e2e_steps, 1):
        parts.append(f"| {CHECKBOX} | {step} | {role_name} | {result} | {i} |")
    parts.extend(
        [
            "",
            "---",
            "",
            "# ملحق: نموذج تقرير الاختبار (يُملأ بعد كل دور)",
            "",
            "```",
            *report_template_lines(),
            "```",
            "",
            "---",
            "",
            f"*آخر تحديث: {issue_date_str()} — نظام مرحّاب*",
        ]
    )
    return "\n".join(parts)


def ar(text: str) -> str:
    try:
        import arabic_reshaper
        from bidi.algorithm import get_display

        return get_display(arabic_reshaper.reshape(text))
    except Exception:
        return text


def build_docx(roles: list[dict]) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.right_margin = Cm(2)
    section.left_margin = Cm(2)

    def set_rtl(paragraph):
        p_pr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)

    def shade_cell(cell, fill: str):
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    def style_run(paragraph, *, size: int, bold: bool = False, color: str | None = None):
        if not paragraph.runs:
            paragraph.add_run(paragraph.text)
        for run in paragraph.runs:
            run.font.size = Pt(size)
            run.bold = bold
            if color:
                run.font.color.rgb = RGBColor.from_string(color)

    # غلاف عصري
    cover = doc.add_table(rows=4, cols=1)
    cover.style = "Table Grid"
    cover_rows = [
        ("دليل اختبار الأدوار والصلاحيات", COVER_TITLE_PT, True, PRIMARY),
        ("نظام مرحّاب", COVER_SUBTITLE_PT, True, ACCENT),
        (f"تاريخ الإصدار: {issue_date_str()}", COVER_META_PT, False, "333355"),
        ("اختبار قبول المستخدم — UAT", COVER_META_PT, False, "555577"),
    ]
    for row_obj, (text, sz, bold, clr) in zip(cover.rows, cover_rows):
        cell = row_obj.cells[0]
        cell.text = text
        shade_cell(cell, HEADER_BG)
        for par in cell.paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_run(par, size=sz, bold=bold, color=clr)
            set_rtl(par)
    doc.add_paragraph("")
    usage = doc.add_paragraph(
        "طريقة الاستخدام: نفّذ كل صف في الجدول وضع علامة ✓ في عمود المربعات عند النجاح "
        "(المربعات قصاد الميزات فقط — وليس قصاد العناوين)."
    )
    set_rtl(usage)

    for role in roles:
        doc.add_page_break()
        h = doc.add_heading(f"{role['num']}. {role['title']} ({role['key']})", level=1)
        set_rtl(h)

        doc.add_heading("نبذة عن الدور", level=2)
        p = doc.add_paragraph(role["intro"])
        set_rtl(p)

        doc.add_heading("نقاط الاختبار — الصلاحيات", level=2)
        counter = 1
        for sec in role["sections"]:
            doc.add_heading(sec["category"], level=3)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = TABLE_HDR_CHECK
            hdr[1].text = "الميزة / نقطة الاختبار"
            hdr[2].text = "#"
            for i, c in enumerate(hdr):
                shade_cell(c, HEADER_BG)
                for par in c.paragraphs:
                    par.runs[0].bold = True
                    if i != 0:
                        par.runs[0].font.size = Pt(BODY_FONT_PT)
                    set_rtl(par)
            for i, item in enumerate(sec["items"]):
                row = table.add_row().cells
                row[0].text = CHECKBOX
                row[1].text = item
                row[2].text = str(counter)
                if i % 2 == 1:
                    for c in row:
                        shade_cell(c, ALT_ROW)
                for ci, c in enumerate(row):
                    for par in c.paragraphs:
                        if ci == 0 and par.runs:
                            par.runs[0].font.size = Pt(CHECKBOX_FONT_PT)
                        set_rtl(par)
                counter += 1
            doc.add_paragraph("")

        doc.add_heading("التقرير المتوقع بعد الاختبار", level=2)
        doc.add_paragraph("بعد اكتمال الاختبار، يجب أن يكون المستخدم قادراً على:")
        for r in role["report"]:
            bp = doc.add_paragraph(r, style="List Bullet")
            set_rtl(bp)
        if role["success"]:
            sp = doc.add_paragraph(f"معايير النجاح: {role['success']}")
            sp.runs[0].bold = True
            set_rtl(sp)

    doc.add_page_break()
    doc.add_heading("ملحق: نموذج تقرير الاختبار", level=1)
    for line in report_template_lines():
        p = doc.add_paragraph(line)
        set_rtl(p)
        if "البيئة:" in line or "□" in line:
            style_run(p, size=CHECKBOX_ENV_FONT_PT)
        elif line.startswith(f"أهم {REPORT_DEFECT_COUNT}"):
            style_run(p, size=APPENDIX_FONT_PT, bold=True)
        else:
            style_run(p, size=APPENDIX_FONT_PT)

    doc.save(DOCX_PATH)


def build_pdf(roles: list[dict]) -> None:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import (
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    font_paths = [
        Path(r"C:\Windows\Fonts\tahoma.ttf"),
        Path(r"C:\Windows\Fonts\arial.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]
    font_name = "Tahoma"
    registered = False
    for fp in font_paths:
        if fp.exists():
            pdfmetrics.registerFont(TTFont(font_name, str(fp)))
            registered = True
            break
    if not registered:
        font_name = "Helvetica"

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleAr",
        parent=styles["Title"],
        fontName=font_name,
        fontSize=18,
        alignment=TA_CENTER,
        textColor=colors.HexColor(f"#{PRIMARY}"),
        spaceAfter=12,
    )
    h1 = ParagraphStyle(
        "H1Ar",
        parent=styles["Heading1"],
        fontName=font_name,
        fontSize=14,
        alignment=TA_RIGHT,
        textColor=colors.HexColor(f"#{PRIMARY}"),
        spaceBefore=14,
        spaceAfter=8,
    )
    h2 = ParagraphStyle(
        "H2Ar",
        parent=styles["Heading2"],
        fontName=font_name,
        fontSize=12,
        alignment=TA_RIGHT,
        textColor=colors.HexColor("#333355"),
        spaceBefore=10,
        spaceAfter=6,
    )
    body = ParagraphStyle(
        "BodyAr",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=10,
        alignment=TA_RIGHT,
        leading=14,
    )
    cell_style = ParagraphStyle(
        "CellAr",
        parent=body,
        fontSize=9,
        leading=12,
    )
    checkbox_style = ParagraphStyle(
        "CheckboxAr",
        parent=cell_style,
        fontSize=CHECKBOX_FONT_PT,
        alignment=TA_CENTER,
        leading=CHECKBOX_FONT_PT + 2,
    )

    purple = colors.HexColor(f"#{HEADER_BG}")
    primary = colors.HexColor(f"#{PRIMARY}")

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        rightMargin=1.5 * cm,
        leftMargin=1.5 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.5 * cm,
        title="Merhab UAT Roles",
    )
    story = []
    cover_data = [
        [Paragraph(ar("دليل اختبار الأدوار والصلاحيات"), title_style)],
        [Paragraph(ar("نظام مرحّاب"), body)],
        [Paragraph(ar(f"تاريخ الإصدار: {date.today().isoformat()}"), body)],
        [Paragraph(ar("UAT — اختبار قبول المستخدم"), body)],
    ]
    cover = Table(cover_data, colWidths=[17 * cm])
    cover.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{HEADER_BG}")),
                ("BOX", (0, 0), (-1, -1), 2, primary),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 18),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 18),
            ]
        )
    )
    story.append(cover)
    story.append(Spacer(1, 0.5 * cm))
    story.append(
        Paragraph(
            ar("طريقة الاستخدام: نفّذ كل صف في الجدول وضع علامة ✓ في عمود «تم» عند النجاح."),
            body,
        )
    )

    for role in roles:
        story.append(PageBreak())
        story.append(Paragraph(ar(f"{role['num']}. {role['title']} ({role['key']})"), h1))
        story.append(Paragraph(ar("نبذة عن الدور"), h2))
        story.append(Paragraph(ar(role["intro"]), body))
        story.append(Spacer(1, 0.2 * cm))
        story.append(Paragraph(ar("نقاط الاختبار — الصلاحيات"), h2))

        counter = 1
        for sec in role["sections"]:
            story.append(Paragraph(ar(sec["category"]), h2))
            data = [
                [
                    Paragraph(CHECKBOX, checkbox_style),
                    Paragraph(ar("الميزة / نقطة الاختبار"), cell_style),
                    Paragraph(ar("#"), cell_style),
                ]
            ]
            for item in sec["items"]:
                data.append(
                    [
                        Paragraph(CHECKBOX, checkbox_style),
                        Paragraph(ar(item), cell_style),
                        Paragraph(str(counter), cell_style),
                    ]
                )
                counter += 1
            col_widths = [1.8 * cm, 13.7 * cm, 1.5 * cm]
            t = Table(data, colWidths=col_widths, repeatRows=1)
            t.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), purple),
                        ("TEXTCOLOR", (0, 0), (-1, 0), primary),
                        ("FONTNAME", (0, 0), (-1, -1), font_name),
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCDD")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("ALIGN", (0, 0), (0, -1), "CENTER"),
                        ("ALIGN", (2, 0), (2, -1), "CENTER"),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor(f"#{ALT_ROW}")]),
                        ("TOPPADDING", (0, 0), (-1, -1), 6),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ]
                )
            )
            story.append(t)
            story.append(Spacer(1, 0.25 * cm))

        story.append(Paragraph(ar("التقرير المتوقع بعد الاختبار"), h2))
        story.append(Paragraph(ar("بعد اكتمال الاختبار، يجب أن يكون المستخدم قادراً على:"), body))
        for r in role["report"]:
            story.append(Paragraph(ar(f"• {r}"), body))
        if role["success"]:
            story.append(Spacer(1, 0.15 * cm))
            story.append(Paragraph(ar(f"معايير النجاح: {role['success']}"), body))

    story.append(PageBreak())
    story.append(Paragraph(ar("ملحق: نموذج تقرير الاختبار"), h1))
    for line in [
        "الدور: _______________",
        "المختبِر: _______________",
        "التاريخ: _______________",
        "البيئة: □ تطوير  □ إنتاج",
        "",
        "عدد النقاط المنفّذة: ___ / ___",
        "عدد النقاط الناجحة:  ___ / ___",
        "عدد العيوب المكتشفة: ___",
        "",
        "أهم 3 عيوب:",
        "1. ",
        "2. ",
        "3. ",
        "",
        "النتيجة النهائية: □ ناجح  □ ناجح بتحفظات  □ فاشل",
        "",
        "ملاحظات:",
    ]:
        story.append(Paragraph(ar(line) if line else " ", body))

    doc.build(story)


def main():
    source_path = SOURCE_PATH if SOURCE_PATH.exists() else MD_PATH
    source = source_path.read_text(encoding="utf-8")
    roles = parse_markdown_roles(source)
    if not roles:
        raise SystemExit(f"تعذّر تحليل الأدوار من {source_path}")
    total_items = sum(len(i) for r in roles for s in r["sections"] for i in [s["items"]])
    if total_items < 50:
        raise SystemExit(
            f"عدد نقاط الاختبار قليل جداً ({total_items}). "
            f"تأكد من وجود {_uat_source_name()} بنقاط - [ ]."
        )

    new_md = build_markdown(roles)
    MD_PATH.write_text(new_md, encoding="utf-8")
    print(f"Markdown: {MD_PATH}")

    build_docx(roles)
    print(f"Word:     {DOCX_PATH}")

    build_pdf(roles)
    print(f"PDF:      {PDF_PATH}")


def _uat_source_name() -> str:
    return SOURCE_PATH.name


if __name__ == "__main__":
    main()

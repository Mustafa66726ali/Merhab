#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""توليد دليل تشغيل الفعاليات — إجراءات المستخدم والحقول فقط."""

from __future__ import annotations

from datetime import date
from pathlib import Path

DOCS = Path(__file__).resolve().parent
MD_PATH = DOCS / "دليل-تشغيل-الفعاليات.md"
DOCX_PATH = DOCS / "دليل-تشغيل-الفعاليات.docx"
PDF_PATH = DOCS / "دليل-تشغيل-الفعاليات.pdf"

PRIMARY = "5B2EFF"
HEADER_BG = "EDE9FF"
ALT_ROW = "F8F7FC"
ACCENT = "7C5CFF"
BODY_FONT_PT = 10
COVER_TITLE_PT = 26
COVER_SUBTITLE_PT = 15

# ── مخطط الرحلة (جداول نصية — بدون رموز Unicode للرسوم) ─────────────────────

JOURNEY_ROWS = [
    ("١", "إنشاء الفعالية", "مدير المنصة / مدير الفعالية", "فعالية «مسودة»"),
    ("٢", "هيكلة: أقسام، مجموعات، جدول", "مدير الفعالية", "تنظيم الضيوف والبرنامج"),
    ("٣", "إضافة الضيوف", "مدير الفعالية / منظم", "ضيوف بحالة «جديد»"),
    ("٤", "إرسال الدعوات", "مدير الفعالية / منظم", "ضيوف «مدعو»"),
    ("٥", "تأكيد الحضور والتذكير", "الضيف + المنظم", "مؤكد / معتذر + QR"),
    ("٦", "تجهيز الفريق والمقاعد والبث", "مدير الفعالية", "فريق + مخطط مقاعد"),
    ("٧", "يوم الحفل", "الجميع", "حضر ← جلس ← بث"),
    ("٨", "التقارير والإغلاق", "مدير المنصة / مدير الفعالية", "فعالية «منتهية»"),
]

GUEST_STATUS_ROWS = [
    ("جديد", "أُضيف للقائمة ولم تُرسل له دعوة بعد"),
    ("مدعو", "وصلته الدعوة بنجاح"),
    ("مؤكد الحضور", "ضغط «تأكيد» أو «نعم» — يظهر QR"),
    ("معتذر", "ضغط «اعتذار» أو «لا»"),
    ("حضر", "سُجّل دخوله عند البوابة (مسح QR)"),
    ("جلس في مقعده", "أُجلِس على طاولة محددة"),
]

GUEST_FLOW_ROWS = [
    ("جديد", "إرسال دعوة", "مدعو"),
    ("مدعو", "تأكيد الحضور", "مؤكد + QR"),
    ("مدعو", "اعتذار", "معتذر"),
    ("مؤكد", "مسح QR عند البوابة", "حضر"),
    ("حضر", "مسح QR عند الطاولة", "جلس"),
]

ROLES_BRIEF = [
    ("مدير المنصة", "ينشئ الفعاليات ويدير أعضاء المنصة والضيوف على مستوى المنصة"),
    ("مدير الفعالية", "يملك الفعالية: هيكلها، ضيوفها، دعواتها، مقاعدها، البث، بدء/إنهاء الحفل"),
    ("منظم الفعالية", "ينفّذ مهام يوم الحفل حسب الصلاحيات: ضيوف، دعوات، رسائل"),
    ("مدير الدخول", "يمسح QR عند البوابة لتحويل الضيف من «مؤكد» إلى «حضر»"),
    ("المنسّق", "يجلس الضيوف على الطاولات بعد تسجيل حضورهم"),
    ("الضيف", "يفتح رابط الدعوة ويؤكد أو يعتذر — لا يملك حساباً في النظام"),
]

# ── المراحل التفصيلية ───────────────────────────────────────────────────────

PHASES = [
    {
        "num": "١",
        "title": "إنشاء الفعالية",
        "role": "مدير المنصة أو مدير الفعالية",
        "nav": "لوحة المنصة ← إدارة المناسبات ← زر «إضافة مناسبة»",
        "sections": [
            {
                "name": "تسجيل الدخول",
                "steps": [
                    "افتح صفحة تسجيل الدخول في النظام.",
                    "أدخل البريد الإلكتروني وكلمة المرور.",
                    "اضغط «تسجيل الدخول».",
                ],
                "fields": [
                    ("البريد الإلكتروني", "مطلوب", "حسابك المسجّل في المنصة", "example@email.com"),
                    ("كلمة المرور", "مطلوب", "كلمة مرور الحساب", "—"),
                ],
                "buttons": ["تسجيل الدخول"],
                "result": "تدخل إلى لوحة التحكم حسب دورك.",
            },
            {
                "name": "نموذج إضافة مناسبة",
                "steps": [
                    "من القائمة الجانبية اختر «إدارة المناسبات».",
                    "اضغط «إضافة مناسبة».",
                    "املأ الحقول أدناه.",
                    "اضغط «حفظ المناسبة».",
                ],
                "fields": [
                    ("اسم المناسبة", "مطلوب", "اسم الحفل كما يظهر في الدعوة والتقارير", "مجلس سعادة الشيخ فهد آل ثاني"),
                    ("الموقع", "اختياري", "اسم القاعة أو المكان النصي", "قاعة الاحتفالات الرئيسية"),
                    ("تاريخ البدء", "مطلوب", "تاريخ يوم الحفل", "2026-07-15"),
                    ("وقت البدء", "مطلوب", "ساعة بدء الحفل", "19:00"),
                    ("وقت الانتهاء", "اختياري", "ساعة انتهاء الحفل (نفس اليوم)", "23:00"),
                    ("مدير الفعالية", "اختياري", "اختر عضواً بدور «مدير فعالية» من قائمة المنصة", "—"),
                    ("منظم الفعالية", "اختياري", "اختر عضواً بدور «منظم فعالية»", "—"),
                    ("الموقع الجغرافي (نص)", "اختياري", "عنوان يظهر في الدعوة والخريطة", "قصر الوجبة، الدوحة"),
                    ("الإحداثيات", "اختياري", "انقر على الخريطة أو أدخل خط العرض والطول", "25.2854, 51.5310"),
                    ("صورة الغلاف", "اختياري", "صورة تظهر في صفحة الدعوة (حد أقصى 5 ميغابايت)", "—"),
                ],
                "buttons": ["حفظ المناسبة"],
                "result": "تُنشأ فعالية جديدة بحالة «مسودة» وتظهر في قائمة المناسبات.",
            },
        ],
    },
    {
        "num": "٢",
        "title": "هيكلة الفعالية",
        "role": "مدير الفعالية",
        "nav": "لوحة مدير الفعالية ← اختر المناسبة",
        "sections": [
            {
                "name": "إضافة قسم",
                "steps": [
                    "افتح صفحة تفاصيل المناسبة.",
                    "في قسم «الأقسام والمجموعات» اضغط «إضافة قسم».",
                    "املأ الحقول واضغط «حفظ».",
                ],
                "fields": [
                    ("اسم القسم", "مطلوب", "تصنيف الضيوف", "كبار الشخصيات"),
                    ("موقع القسم", "اختياري", "مكان القسم داخل الموقع", "الطابق الأرضي - القاعة 4"),
                    ("لون القسم", "اختياري", "لون يميّز القسم في الواجهة", "اختر من منتقي الألوان"),
                    ("وصف", "اختياري", "ملاحظات عن القسم", "—"),
                ],
                "buttons": ["حفظ"],
                "result": "يظهر القسم ويمكن ربط الضيوف به لاحقاً.",
            },
            {
                "name": "إضافة مجموعة داخل قسم",
                "steps": [
                    "بجانب القسم المطلوب اضغط «إضافة مجموعة».",
                    "أدخل موقع المجموعة — يُولَّد الاسم تلقائياً من اسم القسم + الموقع.",
                    "اضغط «حفظ».",
                ],
                "fields": [
                    ("موقع المجموعة", "مطلوب", "مكان المجموعة داخل القسم", "الطابق الأرضي - القاعة 4"),
                ],
                "buttons": ["حفظ"],
                "result": "تُنشأ مجموعة باسم مثل: «كبار الشخصيات - الطابق الأرضي - القاعة 4».",
            },
            {
                "name": "إضافة نشاط في الجدول الزمني",
                "steps": [
                    "من القائمة اختر «الجدول الزمني».",
                    "اضغط «إضافة نشاط» أو أضف صفاً جديداً.",
                    "املأ تفاصيل النشاط واحفظ.",
                ],
                "fields": [
                    ("اسم النشاط", "مطلوب", "عنوان الفقرة", "استقبال الضيوف"),
                    ("وقت البداية", "مطلوب", "ساعة بدء النشاط", "18:30"),
                    ("وقت الانتهاء", "مطلوب", "ساعة انتهاء النشاط", "19:00"),
                    ("الموقع", "اختياري", "«عام» لجميع المواقع أو «موقع مخصص»", "القاعة 4، المسرح الرئيسي"),
                    ("الوصف", "اختياري", "تفاصيل النشاط", "استقبال الضيوف وتقديم المرطبات"),
                ],
                "buttons": ["حفظ"],
                "result": "يظهر النشاط في الجدول ويراه الضيف في صفحة دعوته.",
            },
        ],
    },
    {
        "num": "٣",
        "title": "إضافة الضيوف",
        "role": "مدير الفعالية أو منظم الفعالية (بصلاحية تعديل الضيوف)",
        "nav": "لوحة الفعالية ← الضيوف ← «إضافة ضيف»",
        "sections": [
            {
                "name": "إضافة ضيف جديد",
                "steps": [
                    "اضغط «إضافة ضيف».",
                    "اختر «ضيف جديد».",
                    "املأ البيانات واختر القسم والمجموعة إن وُجدت.",
                    "اضغط «إضافة الضيف».",
                ],
                "fields": [
                    ("الاسم الكامل", "مطلوب", "الاسم كما يظهر في الدعوة", "محمد أحمد العلي"),
                    ("البريد الإلكتروني", "اختياري", "لا يتكرر داخل نفس الفعالية", "guest@email.com"),
                    ("رقم الجوال", "مستحسن", "لإرسال الدعوة عبر واتساب — لا يتكرر داخل الفعالية", "5xxxxxxxx"),
                    ("القسم", "اختياري", "اختر من الأقسام المُنشأة", "كبار الشخصيات"),
                    ("المجموعة", "اختياري", "تظهر بعد اختيار القسم", "—"),
                ],
                "buttons": ["إضافة الضيف"],
                "result": "يُضاف الضيف بحالة «جديد» — لم تُرسل له دعوة بعد.",
            },
            {
                "name": "إضافة ضيف من مناسبة سابقة",
                "steps": [
                    "في نافذة «إضافة ضيف» اختر «من مناسبة سابقة».",
                    "ابحث بالاسم أو الجوال أو البريد.",
                    "اختر الضيف من القائمة.",
                    "اضغط «إضافة الضيف».",
                ],
                "fields": [
                    ("بحث", "اختياري", "تصفية دليل الضيوف", "اسم أو رقم"),
                ],
                "buttons": ["إضافة الضيف"],
                "result": "يُنسخ الضيف للفعالية الحالية بدعوة ورمز QR مستقلين.",
            },
            {
                "name": "تعديل أو حذف ضيف",
                "steps": [
                    "من قائمة الضيوف اضغط على الضيف أو أيقونة «تعديل».",
                    "عدّل الحقول المطلوبة واحفظ.",
                    "للحذف: اضغط «حذف» وأكّد.",
                ],
                "fields": [
                    ("الاسم / البريد / الجوال / القسم / المجموعة", "حسب الحاجة", "قابلة للتعديل في حالة «مسودة» أو قبل الإرسال", "—"),
                ],
                "buttons": ["حفظ", "حذف"],
                "result": "تُحدَّث بيانات الضيف في القائمة.",
            },
        ],
    },
    {
        "num": "٤",
        "title": "إرسال الدعوات",
        "role": "مدير الفعالية أو منظم (بصلاحية إرسال الرسائل)",
        "nav": "لوحة الفعالية ← الدعوات",
        "sections": [
            {
                "name": "تحرير نص الدعوة",
                "steps": [
                    "افتح صفحة «الدعوات» للفعالية.",
                    "تأكد أن وضع العمل «دعوة جديدة» (وليس تذكير).",
                    "حرّر حقول القالب — الاسم والفعالية والتاريخ تُدرج تلقائياً.",
                    "اضغط «حفظ القالب».",
                ],
                "fields": [
                    ("عنوان الدعوة", "اختياري", "يظهر أعلى صفحة الضيف", "دعوة خاصة"),
                    ("التحية", "مطلوب", "تسبق اسم الضيف تلقائياً", "مرحباً"),
                    ("سطر المناسبة", "مطلوب", "يتبعه اسم الفعالية تلقائياً", "يسعدنا دعوتك لحضور"),
                    ("إظهار التاريخ والوقت", "تبديل", "يعرض تاريخ ووقت الفعالية في الرسالة", "مفعّل / معطّل"),
                    ("إظهار المكان", "تبديل", "يعرض موقع الفعالية في الرسالة", "مفعّل / معطّل"),
                    ("سطر دعوة الرابط", "مطلوب", "النص قبل رابط الدعوة", "يرجى تأكيد حضورك عبر الرابط التالي:"),
                ],
                "buttons": ["حفظ القالب"],
                "result": "يُحفظ نص الدعوة الافتراضي للفعالية.",
            },
            {
                "name": "اختيار الجمهور والإرسال",
                "steps": [
                    "في قسم «الجمهور» اختر من سيستلم الدعوة.",
                    "راجع المعاينة أسفل الصفحة.",
                    "اضغط «إرسال الدعوات».",
                    "تابع نتائج الإرسال (تم / فشل) لكل ضيف.",
                ],
                "fields": [
                    ("الجمهور: الكل", "خيار", "كل ضيوف الفعالية", "—"),
                    ("الجمهور: قسم", "خيار", "اختر قسماً محدداً", "كبار الشخصيات"),
                    ("الجمهور: مجموعة", "خيار", "اختر مجموعة محددة", "—"),
                    ("الجمهور: ضيوف محددون", "خيار", "حدّد أسماء من القائمة", "—"),
                ],
                "buttons": ["إرسال الدعوات"],
                "result": "يتحول كل ضيف نجح إرساله إلى «مدعو». من فشل الإرسال يبقى «جديد».",
            },
        ],
    },
    {
        "num": "٥",
        "title": "تأكيد الحضور والتذكيرات",
        "role": "الضيف + مدير/منظم الفعالية",
        "nav": "رابط الدعوة للضيف / صفحة الدعوات للمنظم",
        "sections": [
            {
                "name": "ما يفعله الضيف",
                "steps": [
                    "يفتح رابط الدعوة الواصل عبر واتساب أو الرسالة.",
                    "يراجع تفاصيل المناسبة: التاريخ، الوقت، الموقع، الجدول.",
                    "في قسم «هل ستشرفنا بحضورك؟» يختار أحد الخيارين.",
                ],
                "fields": [],
                "buttons": ["تأكيد الحضور", "اعتذار"],
                "result": "تأكيد ← حالة «مؤكد» + ظهور رمز QR. اعتذار ← حالة «معتذر».",
            },
            {
                "name": "تأكيد عبر واتساب (أزرار الرسالة)",
                "steps": [
                    "يستلم الضيف رسالة تفاعلية فيها أزرار.",
                    "يضغط «نعم» للتأكيد أو «لا» للاعتذار.",
                    "يمكنه أيضاً الضغط «فتح الدعوة» لعرض الصفحة الكاملة.",
                ],
                "fields": [],
                "buttons": ["نعم", "لا", "فتح الدعوة", "عرض الخريطة"],
                "result": "نفس نتيجة التأكيد من الصفحة — يُرسل QR للمؤكد.",
            },
            {
                "name": "إرسال تذكير (المنظم)",
                "steps": [
                    "في صفحة «الدعوات» بدّل الوضع إلى «تذكير».",
                    "حرّر نص التذكير لغير المؤكدين وللمؤكدين (قسمان منفصلان).",
                    "اختر الجمهور واضغط «إرسال التذكير».",
                ],
                "fields": [
                    ("تذكير غير المؤكدين — التحية", "مطلوب", "نص التحية", "مرحباً"),
                    ("تذكير غير المؤكدين — سطر المناسبة", "مطلوب", "يتبعه اسم الفعالية", "يسعدنا دعوتك لحضور"),
                    ("تذكير المؤكدين — التحية", "مطلوب", "عنوان التذكير", "تذكير بموعد المناسبة"),
                    ("تذكير المؤكدين — سطر الدعوة", "مطلوب", "قبل رابط QR", "احتفظ ببطاقة دخولك (QR):"),
                    ("التذكير التلقائي قبل الحفل", "تبديل", "يُرسل تلقائياً قبل الموعد بعدد ساعات محدد", "مثلاً 3 ساعات"),
                ],
                "buttons": ["إرسال التذكير", "حفظ إعدادات التذكير"],
                "result": "غير المؤكدين يستلمون دعوة تفاعلية بعنوان «تذكير». المؤكدون يستلمون نصاً بموعد الحفل.",
            },
        ],
    },
    {
        "num": "٦",
        "title": "تجهيز يوم الحفل",
        "role": "مدير الفعالية",
        "nav": "لوحة الفعالية ← فريق العمل / توزيع المقاعد / البث",
        "sections": [
            {
                "name": "إضافة منسّق أو مدير دخول",
                "steps": [
                    "من القائمة اختر «المنسقون ومدراء الدخول».",
                    "اضغط «إضافة عضو».",
                    "املأ بيانات الحساب واختر الدور.",
                    "بعد الإنشاء: اضغط «تعيين على فعالية» واختر المناسبة.",
                ],
                "fields": [
                    ("الاسم الأول", "اختياري", "—", "أحمد"),
                    ("اسم العائلة", "اختياري", "—", "العلي"),
                    ("البريد الإلكتروني", "مطلوب", "يُستخدم لتسجيل الدخول", "staff@email.com"),
                    ("كلمة المرور", "مطلوب", "6 أحرف على الأقل", "—"),
                    ("الدور", "مطلوب", "منسّق أو مدير دخول", "منسّق / مدير دخول"),
                    ("نوع المنسق", "مطلوب للمنسق", "وصف يظهر للضيف", "منسق رجال"),
                ],
                "buttons": ["إنشاء الحساب", "تعيين على فعالية"],
                "result": "لا يصل العضو لصلاحيات المسح أو الإجلاس إلا بعد التعيين على الفعالية.",
            },
            {
                "name": "إنشاء مخطط مقاعد وطاولة",
                "steps": [
                    "افتح «توزيع المقاعد».",
                    "اضغط «مخطط جديد» وأدخل اسمه.",
                    "اضغط «إضافة طاولة» وحدد الشكل وعدد الكراسي.",
                    "اسحب الطاولات على المخطط أو استخدم المحرّر البصري.",
                    "احفظ المخطط.",
                ],
                "fields": [
                    ("اسم المخطط", "مطلوب", "اسم القاعة أو الصالة", "القاعة الرئيسية"),
                    ("وصف المخطط", "اختياري", "ملاحظات", "—"),
                    ("اسم الطاولة", "مطلوب", "رقم أو اسم الطاولة", "طاولة 1"),
                    ("عدد الكراسي", "مطلوب", "سعة الجلوس", "8"),
                    ("شكل الطاولة", "مطلوب", "دائري / مستطيل / مربع", "—"),
                    ("القسم", "اختياري", "ربط الطاولة بقسم", "—"),
                    ("المجموعة", "اختياري", "طاولة واحدة لكل مجموعة", "—"),
                ],
                "buttons": ["حفظ المخطط", "إضافة طاولة"],
                "result": "مخطط جاهز لإجلاس الضيوف يوم الحفل.",
            },
            {
                "name": "إعداد البث المباشر",
                "steps": [
                    "افتح «البث المباشر».",
                    "فعّل «تفعيل البث للضيوف».",
                    "اختر نوع البث واملأ الحقل المناسب.",
                    "اضغط «حفظ الإعدادات».",
                ],
                "fields": [
                    ("تفعيل البث للضيوف", "تبديل", "يظهر قسم البث في صفحة دعوة الضيف", "مفعّل"),
                    ("نوع البث: متوقف", "خيار", "لا يظهر شيء للضيوف", "—"),
                    ("نوع البث: ملف صوتي", "خيار", "ارفع MP3 أو M4A أو WAV", "—"),
                    ("نوع البث: يوتيوب", "خيار", "رابط فيديو أو بث مباشر", "https://youtube.com/..."),
                    ("نوع البث: ميكروفون", "خيار", "بث صوتي مباشر من الجهاز", "—"),
                    ("نوع البث: كاميرا", "خيار", "بث فيديو مباشر من الكاميرا", "—"),
                ],
                "buttons": ["حفظ الإعدادات", "بدء البث", "إيقاف البث"],
                "result": "يظهر البث للضيوف في صفحة دعوتهم عند التفعيل.",
            },
        ],
    },
    {
        "num": "٧",
        "title": "يوم الحفل — التشغيل",
        "role": "مدير الفعالية + مدير الدخول + المنسّق + الضيوف",
        "nav": "صفحة تفاصيل المناسبة / بوابة الدخول / الإجلاس",
        "sections": [
            {
                "name": "بدء المناسبة (مدير الفعالية)",
                "steps": [
                    "افتح صفحة تفاصيل المناسبة.",
                    "اضغط الزر الأخضر «بدء المناسبة».",
                    "تتحول حالة الفعالية إلى «تعمل الآن».",
                ],
                "fields": [],
                "buttons": ["بدء المناسبة"],
                "result": "يُفعَّل تسجيل الحضور والإجلاس. قبل البدء تظهر رسالة أن هذه العمليات غير متاحة.",
            },
            {
                "name": "تسجيل دخول الضيف (مدير الدخول)",
                "steps": [
                    "سجّل الدخول بحساب «مدير دخول».",
                    "من القائمة اختر «تسجيل الدخول».",
                    "اضغط «مسح بالكاميرا» أو أدخل رمز QR يدوياً.",
                    "وجّه الكاميرا نحو QR الضيف عند البوابة.",
                ],
                "fields": [
                    ("رمز QR", "مطلوب", "يظهر في دعوة الضيف بعد التأكيد", "—"),
                ],
                "buttons": ["مسح بالكاميرا"],
                "result": "يتحول الضيف من «مؤكد» إلى «حضر» — يجب أن تكون الفعالية «تعمل الآن».",
            },
            {
                "name": "إجلاس الضيف (المنسّق)",
                "steps": [
                    "سجّل الدخول بحساب «منسّق».",
                    "افتح «إجلاس الضيوف» واختر المناسبة.",
                    "فعّل «وضع المسح» أو انقر على كرسي فارغ.",
                    "امسح QR الضيف لإجلاسه على المقعد.",
                ],
                "fields": [
                    ("المقعد / الكرسي", "مطلوب", "كرسي فارغ على المخطط", "—"),
                    ("رمز QR", "مطلوب", "QR الضيف", "—"),
                ],
                "buttons": ["مسح", "إلغاء الإجلاس"],
                "result": "يتحول الضيف من «حضر» إلى «جلس في مقعده». لا يُجلَس من لم يُسجَّل حضوره.",
            },
            {
                "name": "إرسال رابط البث (مدير الفعالية)",
                "steps": [
                    "افتح «البث المباشر».",
                    "ابدأ البث إن كان ميكروفوناً أو كاميرا.",
                    "اضغط «إرسال رابط المشاهدة للضيوف».",
                    "يصل الرابط للحاضرين عبر واتساب.",
                ],
                "fields": [],
                "buttons": ["بدء البث", "إرسال رابط المشاهدة", "نسخ الرابط"],
                "result": "يضغط الضيف «مشاهدة» في الرسالة لفتح صفحة البث.",
            },
            {
                "name": "إنهاء المناسبة (مدير الفعالية)",
                "steps": [
                    "بعد انتهاء الحفل افتح صفحة تفاصيل المناسبة.",
                    "اضغط «إنهاء المناسبة» وأكّد.",
                ],
                "fields": [],
                "buttons": ["إنهاء المناسبة"],
                "result": "تتحول الحالة إلى «منتهية» — لا يُسمح بتسجيل حضور أو إجلاس جديد.",
            },
        ],
    },
    {
        "num": "٨",
        "title": "بعد انتهاء الحفل",
        "role": "مدير المنصة / مدير الفعالية",
        "nav": "لوحة المنصة ← التقارير",
        "sections": [
            {
                "name": "مراجعة التقارير",
                "steps": [
                    "افتح «التقارير والإحصائيات».",
                    "اختر الفعالية من القائمة.",
                    "راجع: عدد المدعوين، المؤكدين، الحاضرين، المعتذرين.",
                    "صدّر تقرير التهنئات PDF إن رغبت.",
                ],
                "fields": [
                    ("الفعالية", "مطلوب", "اختر المناسبة المطلوبة", "—"),
                ],
                "buttons": ["تصدير PDF"],
                "result": "تقرير نهائي عن أداء الفعالية.",
            },
            {
                "name": "أرشفة وإعادة استخدام الضيوف",
                "steps": [
                    "من قائمة المناسبات يمكنك أرشفة الفعالية المنتهية.",
                    "عند إنشاء فعالية جديدة استخدم «من مناسبة سابقة» لإضافة ضيوف سابقين بسرعة.",
                ],
                "fields": [],
                "buttons": ["أرشفة"],
                "result": "بيانات الضيوف تبقى في دليل المنصة للفعاليات القادمة.",
            },
        ],
    },
]


def issue_date_str() -> str:
    return date.today().strftime("%d-%m-%Y")


def ar(text: str) -> str:
    try:
        import arabic_reshaper
        from bidi.algorithm import get_display

        return get_display(arabic_reshaper.reshape(text))
    except Exception:
        return text


def _md_table(headers: list[str], rows: list[tuple]) -> list[str]:
    lines = [
        "| " + " | ".join(headers) + " |",
        "|" + "|".join("---" for _ in headers) + "|",
    ]
    for row in rows:
        lines.append("| " + " | ".join(str(c) for c in row) + " |")
    return lines


def build_markdown() -> str:
    lines = [
        "# دليل تشغيل الفعاليات — نظام مرحّاب",
        "",
        f"**تاريخ الإصدار:** {issue_date_str()}",
        "",
        "> دليل عملي يصف **ماذا يفعل كل مستخدم** خطوة بخطوة، مع تفاصيل الحقول في كل نموذج.",
        "",
        "---",
        "",
        "## مخطط الرحلة",
        "",
    ]
    lines.extend(_md_table(["الخطوة", "المرحلة", "من ينفّذ", "النتيجة"], JOURNEY_ROWS))
    lines.extend(["", "## الأدوار باختصار", ""])
    lines.extend(_md_table(["الدور", "ماذا يفعل"], ROLES_BRIEF))
    lines.extend(["", "## حالات الضيف", ""])
    lines.extend(_md_table(["الحالة", "المعنى"], GUEST_STATUS_ROWS))
    lines.extend(["", "## انتقالات حالة الضيف", ""])
    lines.extend(_md_table(["من", "الإجراء", "إلى"], GUEST_FLOW_ROWS))
    lines.extend(["", "---", ""])

    for phase in PHASES:
        lines.extend([
            f"## المرحلة {phase['num']}: {phase['title']}",
            "",
            f"**من ينفّذ:** {phase['role']}  ",
            f"**أين في النظام:** {phase['nav']}",
            "",
        ])
        for sec in phase["sections"]:
            lines.extend([f"### {sec['name']}", ""])
            lines.append("**الخطوات:**")
            for i, step in enumerate(sec["steps"], 1):
                lines.append(f"{i}. {step}")
            if sec["fields"]:
                lines.extend(["", "**الحقول:**", ""])
                lines.extend(_md_table(
                    ["الحقل", "الإلزام", "الوصف", "مثال"],
                    sec["fields"],
                ))
            if sec.get("buttons"):
                lines.extend(["", "**الأزرار:** " + " — ".join(sec["buttons"])])
            lines.extend(["", f"**النتيجة:** {sec['result']}", ""])
        lines.append("---")
        lines.append("")

    lines.append(f"*آخر تحديث: {issue_date_str()} — نظام مرحّاب*")
    return "\n".join(lines)


# ── Word ──────────────────────────────────────────────────────────────────────

def build_docx() -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.right_margin = Cm(2)
    section.left_margin = Cm(2)

    def set_rtl(paragraph):
        p_pr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)

    def shade_cell(cell, fill: str):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    def add_heading_rtl(text: str, level: int):
        h = doc.add_heading(text, level=level)
        set_rtl(h)
        return h

    def add_para_rtl(text: str, *, bold: bool = False, size: int = BODY_FONT_PT):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        set_rtl(p)
        return p

    def add_table_rtl(headers: list[str], rows: list[tuple], col_widths=None):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        for i, h in enumerate(headers):
            t.rows[0].cells[i].text = h
            shade_cell(t.rows[0].cells[i], HEADER_BG)
            set_rtl(t.rows[0].cells[i].paragraphs[0])
        for idx, row in enumerate(rows):
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
                if idx % 2:
                    shade_cell(cells[i], ALT_ROW)
                set_rtl(cells[i].paragraphs[0])
        return t

    # غلاف
    cover = doc.add_table(rows=3, cols=1)
    cover.style = "Table Grid"
    for row_obj, (text, sz, bold, clr) in zip(
        cover.rows,
        [
            ("دليل تشغيل الفعاليات", COVER_TITLE_PT, True, PRIMARY),
            ("نظام مرحّاب — إجراءات المستخدم", COVER_SUBTITLE_PT, True, ACCENT),
            (f"تاريخ الإصدار: {issue_date_str()}", 12, False, "333355"),
        ],
    ):
        cell = row_obj.cells[0]
        cell.text = text
        shade_cell(cell, HEADER_BG)
        for par in cell.paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in par.runs:
                run.font.size = Pt(sz)
                run.bold = bold
                run.font.color.rgb = RGBColor.from_string(clr)
            set_rtl(par)

    doc.add_page_break()
    add_heading_rtl("مخطط الرحلة", 1)
    add_table_rtl(["الخطوة", "المرحلة", "من ينفّذ", "النتيجة"], JOURNEY_ROWS)
    add_heading_rtl("الأدوار", 1)
    add_table_rtl(["الدور", "ماذا يفعل"], ROLES_BRIEF)
    add_heading_rtl("حالات الضيف", 1)
    add_table_rtl(["الحالة", "المعنى"], GUEST_STATUS_ROWS)
    add_heading_rtl("انتقالات حالة الضيف", 2)
    add_table_rtl(["من", "الإجراء", "إلى"], GUEST_FLOW_ROWS)

    for phase in PHASES:
        doc.add_page_break()
        add_heading_rtl(f"المرحلة {phase['num']}: {phase['title']}", 1)
        add_para_rtl(f"من ينفّذ: {phase['role']}", bold=True)
        add_para_rtl(f"أين: {phase['nav']}")
        for sec in phase["sections"]:
            add_heading_rtl(sec["name"], 2)
            add_para_rtl("الخطوات:", bold=True)
            for i, step in enumerate(sec["steps"], 1):
                add_para_rtl(f"{i}. {step}")
            if sec["fields"]:
                add_para_rtl("الحقول:", bold=True)
                add_table_rtl(["الحقل", "الإلزام", "الوصف", "مثال"], sec["fields"])
            if sec.get("buttons"):
                add_para_rtl("الأزرار: " + " — ".join(sec["buttons"]), bold=True)
            add_para_rtl(f"النتيجة: {sec['result']}")

    doc.save(DOCX_PATH)


# ── PDF ───────────────────────────────────────────────────────────────────────

def build_pdf() -> None:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    font_name = "Tahoma"
    for fp in [
        Path(r"C:\Windows\Fonts\tahoma.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]:
        if fp.exists():
            pdfmetrics.registerFont(TTFont(font_name, str(fp)))
            break
    else:
        font_name = "Helvetica"

    title_style = ParagraphStyle(
        "TitleAr", fontName=font_name, fontSize=COVER_TITLE_PT,
        alignment=TA_CENTER, textColor=colors.HexColor(f"#{PRIMARY}"), spaceAfter=8,
    )
    cover_sub = ParagraphStyle(
        "CoverSub", fontName=font_name, fontSize=COVER_SUBTITLE_PT,
        alignment=TA_CENTER, textColor=colors.HexColor(f"#{ACCENT}"),
    )
    h1 = ParagraphStyle(
        "H1", fontName=font_name, fontSize=14, alignment=TA_RIGHT,
        textColor=colors.HexColor(f"#{PRIMARY}"), spaceBefore=12, spaceAfter=6,
    )
    h2 = ParagraphStyle(
        "H2", fontName=font_name, fontSize=11, alignment=TA_RIGHT,
        textColor=colors.HexColor("#333355"), spaceBefore=8, spaceAfter=4,
    )
    body = ParagraphStyle(
        "Body", fontName=font_name, fontSize=BODY_FONT_PT, alignment=TA_RIGHT, leading=14,
    )
    cell_s = ParagraphStyle("Cell", fontName=font_name, fontSize=8, alignment=TA_RIGHT, leading=10)
    header_s = ParagraphStyle(
        "Hdr", fontName=font_name, fontSize=9, alignment=TA_RIGHT,
        textColor=colors.HexColor(f"#{PRIMARY}"), leading=11,
    )
    purple = colors.HexColor(f"#{HEADER_BG}")
    primary = colors.HexColor(f"#{PRIMARY}")

    def pdf_table(headers: list[str], rows: list[tuple], widths: list[float]):
        data = [[Paragraph(ar(h), header_s) for h in headers]]
        for row in rows:
            data.append([Paragraph(ar(str(c)), cell_s) for c in row])
        t = Table(data, colWidths=widths, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), purple),
            ("FONTNAME", (0, 0), (-1, -1), font_name),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCDD")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor(f"#{ALT_ROW}")]),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ]))
        return t

    doc = SimpleDocTemplate(
        str(PDF_PATH), pagesize=A4,
        rightMargin=1.5 * cm, leftMargin=1.5 * cm,
        topMargin=1.5 * cm, bottomMargin=1.5 * cm,
    )
    story = []

    cover = Table(
        [
            [Paragraph(ar("دليل تشغيل الفعاليات"), title_style)],
            [Paragraph(ar("نظام مرحّاب — إجراءات المستخدم"), cover_sub)],
            [Paragraph(ar(f"تاريخ الإصدار: {issue_date_str()}"), cover_sub)],
        ],
        colWidths=[17 * cm],
    )
    cover.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), purple),
        ("BOX", (0, 0), (-1, -1), 2, primary),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 22),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 22),
    ]))
    story.append(cover)
    story.append(PageBreak())

    story.append(Paragraph(ar("مخطط الرحلة"), h1))
    story.append(pdf_table(
        ["الخطوة", "المرحلة", "من ينفّذ", "النتيجة"],
        JOURNEY_ROWS, [1.2 * cm, 5.5 * cm, 4.5 * cm, 5.8 * cm],
    ))
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph(ar("الأدوار"), h1))
    story.append(pdf_table(["الدور", "ماذا يفعل"], ROLES_BRIEF, [4 * cm, 13 * cm]))
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph(ar("حالات الضيف"), h1))
    story.append(pdf_table(["الحالة", "المعنى"], GUEST_STATUS_ROWS, [4 * cm, 13 * cm]))
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph(ar("انتقالات حالة الضيف"), h2))
    story.append(pdf_table(["من", "الإجراء", "إلى"], GUEST_FLOW_ROWS, [3.5 * cm, 6 * cm, 7.5 * cm]))
    story.append(PageBreak())

    for phase in PHASES:
        story.append(Paragraph(ar(f"المرحلة {phase['num']}: {phase['title']}"), h1))
        story.append(Paragraph(ar(f"من ينفّذ: {phase['role']}"), body))
        story.append(Paragraph(ar(f"أين: {phase['nav']}"), body))
        story.append(Spacer(1, 0.15 * cm))
        for sec in phase["sections"]:
            story.append(Paragraph(ar(sec["name"]), h2))
            story.append(Paragraph(ar("الخطوات:"), body))
            for i, step in enumerate(sec["steps"], 1):
                story.append(Paragraph(ar(f"{i}. {step}"), body))
            if sec["fields"]:
                story.append(Spacer(1, 0.1 * cm))
                story.append(Paragraph(ar("الحقول:"), body))
                story.append(pdf_table(
                    ["الحقل", "الإلزام", "الوصف", "مثال"],
                    sec["fields"], [3.2 * cm, 1.8 * cm, 7.5 * cm, 4.5 * cm],
                ))
            if sec.get("buttons"):
                story.append(Paragraph(ar("الأزرار: " + " — ".join(sec["buttons"])), body))
            story.append(Paragraph(ar(f"النتيجة: {sec['result']}"), body))
            story.append(Spacer(1, 0.2 * cm))
        story.append(PageBreak())

    doc.build(story)


def main() -> None:
    MD_PATH.write_text(build_markdown(), encoding="utf-8")
    build_docx()
    build_pdf()
    print("Generated:", MD_PATH.name, DOCX_PATH.name, PDF_PATH.name)


if __name__ == "__main__":
    main()
